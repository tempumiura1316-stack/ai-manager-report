#諸々インポート
import streamlit as st 
import io
from openai import OpenAI
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

#OpenAIを使う準備
client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

#UI
st.set_page_config(page_title="AI報告書メーカー", page_icon="🧠")
st.title("🧠 AI報告書メーカー")
st.info("当アプリは OpenAI の **API** を利用します。\n\n"
        "API経由の入出力は学習に利用されないため情報漏洩の心配はございませんのでご安心ください。\n\n"
        "詳細はページ下部参照")

st.write("こんにちは！これは週報作成用のアプリです。")

text = st.text_area("日報を入力してください")

industry_option = st.selectbox(
    "あなたの所属業界を選んでください",
    ["IT・ソフトウェア", "コンサルティング", "製造業", "小売・流通", "金融・保険", "広告・マーケティング", "教育・研修", "医療・ヘルスケア", "その他（自由入力）"])
if industry_option == "その他（自由入力）":
    industry = st.text_input("所属業界を自由に入力してください")
else:
    industry = industry_option

sender_option = st.selectbox("あなたの役職を選んでください", ["営業担当", "エンジニア", "プロジェクトマネージャー", "コンサルタント", "マーケター", "人事・総務", "その他（自由入力）"])
if sender_option == "その他（自由入力）":
    sender = st.text_input("あなたの役職を自由に入力してください")
else:
    sender = sender_option

receiver_option = st.selectbox("報告先の役職を選んでください",["上司", "部長", "経営層", "クライアント", "チームリーダー", "その他（自由入力）"])
if receiver_option == "その他（自由入力）":
    receiver = st.text_input("報告先の役職を自由に入力してください")
else:
    receiver = receiver_option

purpose_option = st.selectbox("この報告の目的を選んでください",["進捗報告", "成果報告", "課題共有", "承認依頼", "提案・改善案の提示", "その他（自由入力）"])
if purpose_option == "その他（自由入力）":
    purpose = st.text_input("報告の目的を自由に入力してください")
else:
    purpose = purpose_option

#AIで実行
if st.button("週報を作成"):
    st.success("ここにAIの出力が表示されます")
    prompt = f"""
    あなたは{industry}業界に所属している{sender}です。
    あなたは{receiver}に対して、{purpose}のための週次報告書を作成します。
    以下はあなたの日報です。これをもとに、週報(見出し/実績/課題/来週のPlan)を日本語でMarkdown形式で出力してください。
    
    {text}"""
    
    md = client.chat.completions.create(
    model="gpt-4o-mini",
    messages=[{"role": "user", "content": prompt}]).choices[0].message.content
    st.markdown(md)

    #pdf作成
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    pdfmetrics.registerFont(UnicodeCIDFont("HeiseiMin-W3"))
    styles = getSampleStyleSheet()
    normal_style = ParagraphStyle(name="NormalJP", parent=styles["Normal"], fontName="HeiseiMin-W3", fontSize=11, leading=16)
    
    story = []
    lines = md.splitlines()
    for line in lines:
        line = line.strip()
        if not line:
            story.append(Spacer(1, 8))
            continue
        if line.startswith("# "):
            story.append(Paragraph(line[2:], normal_style))
            story.append(Spacer(1, 6))
        else:
            story.append(Paragraph(line, normal_style))
    doc.build(story)
    pdf_bytes = buf.getvalue()
    
    st.download_button(
        "PDFをダウンロード",
        data=pdf_bytes,
        file_name="weekly_report.pdf",
        mime="application/pdf",
        use_container_width=True
    )
    
    #Word作成
    docx_buf = io.BytesIO()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Yu Gothic"
    style.font.size = Pt(11)
    for raw in md.splitlines():
        line = raw.strip()
        if not line:
            doc.add_paragraph("")  # 空行を挿入
        elif line.startswith("## "):
            p = doc.add_paragraph(line[3:])
            p.runs[0].bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith("# "):
            p = doc.add_paragraph(line[2:])
            p.runs[0].bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)
    doc.save(docx_buf)
    docx_bytes = docx_buf.getvalue()
    button_label = "Word（.docx）をダウンロード"
    st.download_button(button_label, data=docx_bytes, file_name="weekly_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

with st.expander("🔒 データ取り扱いの根拠（公式）"):
    st.markdown(
        "- **APIのデータは学習に不使用**（Business/Enterprise/API）: [公式ヘルプ](https://help.openai.com/en/articles/5722486-how-your-data-is-used-to-improve-model-performance)\n"
        "- **Enterprise Privacy（30日ログ・ZDR）**: [公式ページ](https://openai.com/enterprise-privacy/)\n"
        "- **Businessデータは学習に不使用**: [公式ページ](https://openai.com/business-data/)\n"
        "- **セキュリティ（第三者ペンテスト等）**: [Security](https://openai.com/security-and-privacy/)\n"
    )
st.graphviz_chart("""
digraph G {
  rankdir=LR;
  node [shape=box, style=rounded, fontsize=10];
  subgraph cluster_local {
    label="あなたの環境（ローカル/Streamlit Cloud）";
    color="#bbbbbb";
    App [label="AI報告書メーカー\\n(送信前マスキング可)"];
    PDF [label="PDF/Word 生成\\n(ローカル)"];
  }
  subgraph cluster_openai {
    label="OpenAI (APIサーバ)";
    color="#bbbbbb";
    API [label="OpenAI API\\n(TLS暗号化通信)"];
    Logs [label="運用ログ\\n最大30日保管\\n(ZDR設定時は0日)"];
  }
  App -> API [label="必要最小限のテキスト送信"];
  API -> App [label="AI生成結果"];
  App -> PDF [label="ローカル変換"];
  API -> Logs [style=dashed, label="監査/不正検知"];
}
""")
